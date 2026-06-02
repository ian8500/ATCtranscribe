from datetime import datetime
import logging
import os
import tempfile
import time

from fastapi import FastAPI, Depends, HTTPException, status, UploadFile, File, Response, Request, BackgroundTasks
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from starlette.background import BackgroundTask
from sqlalchemy.orm import Session
from sqlalchemy import select
from .db import get_db, ensure_database_ready
from .models import (
    User, Transcript, TranscriptLine, SpeakerLabel, TranscriptStatus, AccessLevel,
    ExcludeEntry, VocabularyEntry, TranscriptionJob, TranscriptionJobStatus,
)
from .schemas import (
    UserListItem, LoginRequest, TranscriptCreate, TranscriptOut, TranscriptUpdate,
    TranscriptLineOut, TranscriptLineCreate, TranscriptLineUpdate, StartTranscriptionRequest,
    SplitLineRequest, MergeLinesRequest, SpeakerLabelCreate, SpeakerLabelOut, SpeakerLabelUpdate,
    ForgotPasswordRequest, UserCreate, UserOut, UserUpdate, ResetPasswordRequest, AuditLogOut,
    EntryCreate, EntryOut, TranscriptionJobOut,
)
from .security import verify_password, create_access_token, hash_password
from .deps import get_current_user, require_admin
from .settings import get_settings
from .audio import save_wav_file, purge_transcript_audio, sanitize_filename
from .transcription import transcribe_audio, hms_to_seconds
from .emailer import send_email
from .audit import log_event
from .rate_limit import rate_limiter
from .models import AuditLog
from docx import Document


settings = get_settings()
logger = logging.getLogger("atc.api")

app = FastAPI(title="ATC Transcriber")

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.allowed_cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)


@app.on_event("startup")
def sweep_completed_audio() -> None:
    settings.validate_security()
    ensure_database_ready()
    from .db import SessionLocal
    db = SessionLocal()
    try:
        completed = db.execute(select(Transcript).where(Transcript.status == TranscriptStatus.completed)).scalars().all()
        for transcript in completed:
            if transcript.wav_storage_path and os.path.exists(transcript.wav_storage_path):
                purge_transcript_audio(transcript.id)
                transcript.wav_storage_path = None
                transcript.wav_filename = None
                log_event(db, None, "audio_purged_on_startup", "transcript", transcript.id)
        db.commit()
    finally:
        db.close()


@app.get("/api/health")
def health():
    return {"status": "ok", "app": "ATC Transcriber"}


@app.get("/api/auth/user-list", response_model=list[UserListItem])
def user_list(db: Session = Depends(get_db)):
    users = db.execute(select(User).where(User.active == True)).scalars().all()
    return [UserListItem(id=user.id, name=user.name) for user in users]


@app.post("/api/auth/login")
def login(payload: LoginRequest, response: Response, request: Request, db: Session = Depends(get_db)):
    rate_limiter.check(f"login:{request.client.host}:{payload.user_id}")
    user = db.get(User, payload.user_id)
    if not user or not user.active:
        log_event(db, None, "login_failed", "user", payload.user_id, {"reason": "inactive_or_missing"})
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid credentials")
    if not verify_password(payload.password, user.password_hash):
        log_event(db, user.id, "login_failed", "user", user.id)
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid credentials")
    token = create_access_token(user.id)
    response.set_cookie(
        settings.session_cookie_name,
        token,
        httponly=True,
        secure=settings.secure_cookies,
        samesite="strict",
        max_age=settings.access_token_expire_minutes * 60,
    )
    log_event(db, user.id, "login_success", "user", user.id)
    return {"message": "ok"}


@app.post("/api/auth/logout")
def logout(response: Response, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    response.delete_cookie(settings.session_cookie_name, secure=settings.secure_cookies, samesite="strict")
    log_event(db, user.id, "logout", "user", user.id)
    return {"message": "logged out"}


@app.get("/api/auth/me", response_model=UserOut)
def me(user: User = Depends(get_current_user)):
    return user


@app.post("/api/auth/forgot-password")
def forgot_password(payload: ForgotPasswordRequest, request: Request, db: Session = Depends(get_db)):
    rate_limiter.check(f"forgot:{request.client.host}:{payload.user_id}")
    user = db.get(User, payload.user_id)
    if not user:
        return {"message": "If the user exists, admins have been notified."}
    admins = db.execute(select(User).where(User.access_level == AccessLevel.admin)).scalars().all()
    if admins:
        send_email(
            "ATC Transcriber password reset request",
            f"User {user.name} requested a password reset.",
            [admin.email for admin in admins],
        )
    log_event(db, user.id, "forgot_password_request", "user", user.id)
    return {"message": "Admins notified"}


@app.get("/api/users", response_model=list[UserOut])
def list_users(admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    return db.execute(select(User)).scalars().all()


@app.post("/api/admin/users", response_model=UserOut)
def create_user(payload: UserCreate, admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    user = User(
        name=payload.name,
        email=payload.email,
        password_hash=hash_password(payload.password),
        access_level=payload.access_level,
        active=True,
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    log_event(db, admin.id, "user_created", "user", user.id)
    return user


@app.patch("/api/admin/users/{user_id}", response_model=UserOut)
def update_user(user_id: int, payload: UserUpdate, admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    user = db.get(User, user_id)
    if not user:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="User not found")
    data = payload.dict(exclude_unset=True)
    for key, value in data.items():
        setattr(user, key, value)
    db.commit()
    db.refresh(user)
    log_event(db, admin.id, "user_updated", "user", user.id)
    return user


@app.post("/api/admin/users/{user_id}/reset-password")
def reset_password(user_id: int, payload: ResetPasswordRequest, admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    user = db.get(User, user_id)
    if not user:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="User not found")
    user.password_hash = hash_password(payload.new_password)
    db.commit()
    send_email(
        "ATC Transcriber password reset",
        f"Your password was reset by an admin. Please log in and change it.",
        [user.email],
    )
    log_event(db, admin.id, "password_reset", "user", user.id)
    return {"message": "password reset"}


@app.get("/api/admin/audit-log", response_model=list[AuditLogOut])
def audit_log(admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    logs = db.execute(select(AuditLog).order_by(AuditLog.created_at.desc()).limit(200)).scalars().all()
    return logs


def get_transcript_or_404(transcript_id: int, user: User, db: Session) -> Transcript:
    transcript = db.get(Transcript, transcript_id)
    if not transcript:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Transcript not found")
    if user.access_level != AccessLevel.admin and transcript.owner_user_id != user.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Forbidden")
    return transcript


def reorder_lines(transcript_id: int, db: Session) -> None:
    lines = db.execute(
        select(TranscriptLine)
        .where(TranscriptLine.transcript_id == transcript_id)
        .order_by(TranscriptLine.order_index, TranscriptLine.id)
    ).scalars().all()
    for idx, line in enumerate(lines):
        line.order_index = idx
    db.commit()


def validate_speaker_label(transcript_id: int, label_id: int | None, db: Session) -> None:
    if label_id is None:
        return
    label = db.get(SpeakerLabel, label_id)
    if not label or label.transcript_id != transcript_id:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Speaker label does not belong to this transcript")


def latest_transcription_job(transcript_id: int, db: Session) -> TranscriptionJob | None:
    return db.execute(
        select(TranscriptionJob)
        .where(TranscriptionJob.transcript_id == transcript_id)
        .order_by(TranscriptionJob.created_at.desc(), TranscriptionJob.id.desc())
        .limit(1)
    ).scalar_one_or_none()


def active_transcription_job(transcript_id: int, db: Session) -> TranscriptionJob | None:
    return db.execute(
        select(TranscriptionJob)
        .where(
            TranscriptionJob.transcript_id == transcript_id,
            TranscriptionJob.status.in_([TranscriptionJobStatus.queued, TranscriptionJobStatus.running]),
        )
        .order_by(TranscriptionJob.created_at.desc(), TranscriptionJob.id.desc())
        .limit(1)
    ).scalar_one_or_none()


def update_job_status(
    db: Session,
    job: TranscriptionJob,
    *,
    status_value: TranscriptionJobStatus | None = None,
    progress: int | None = None,
    error: str | None = None,
) -> None:
    if status_value is not None:
        job.status = status_value
    if progress is not None:
        job.progress = progress
    if error is not None:
        job.error = error
    job.updated_at = datetime.utcnow()
    db.commit()


def run_transcription_job(job_id: int, start_time: str, actor_user_id: int | None = None) -> None:
    from .db import SessionLocal

    db = SessionLocal()
    total_started = time.perf_counter()
    try:
        job = db.get(TranscriptionJob, job_id)
        if not job:
            logger.warning("Transcription job %s was not found", job_id)
            return
        job.status = TranscriptionJobStatus.running
        job.progress = 5
        job.started_at = datetime.utcnow()
        job.updated_at = datetime.utcnow()
        db.commit()

        transcript = db.get(Transcript, job.transcript_id)
        if not transcript:
            raise RuntimeError("Transcript not found")
        if not transcript.wav_storage_path or not os.path.exists(transcript.wav_storage_path):
            raise RuntimeError("No WAV uploaded")

        hms_to_seconds(start_time)
        exclude_entries = db.execute(select(ExcludeEntry).where(ExcludeEntry.transcript_id == transcript.id)).scalars().all()
        vocab_entries = db.execute(select(VocabularyEntry).where(VocabularyEntry.transcript_id == transcript.id)).scalars().all()
        exclude_list = [entry.word_or_phrase for entry in exclude_entries]
        vocab_list = [entry.word_or_phrase for entry in vocab_entries]
        transcript.exclude_snapshot = {"entries": exclude_list}
        transcript.dictionary_snapshot = {"entries": vocab_list}
        update_job_status(db, job, progress=15)

        lines = transcribe_audio(transcript.wav_storage_path, start_time, exclude_list, vocab_list, transcript.id)
        update_job_status(db, job, progress=85)

        write_started = time.perf_counter()
        db.query(TranscriptLine).filter(TranscriptLine.transcript_id == transcript.id).delete()
        for index, line in enumerate(lines):
            db.add(
                TranscriptLine(
                    transcript_id=transcript.id,
                    order_index=index,
                    timestamp_hms=line["timestamp_hms"],
                    text=line["text"],
                    flags_json=line.get("flags_json"),
                )
            )
        transcript.status = TranscriptStatus.in_progress
        job.status = TranscriptionJobStatus.completed
        job.progress = 100
        job.completed_at = datetime.utcnow()
        job.updated_at = datetime.utcnow()
        db.commit()
        logger.info("Transcription database write completed in %.2fs", time.perf_counter() - write_started)
        log_event(db, actor_user_id, "transcription_completed", "transcript", transcript.id, {"job_id": job.id})
        logger.info("Transcription job %s completed in %.2fs", job.id, time.perf_counter() - total_started)
    except Exception as exc:
        db.rollback()
        job = db.get(TranscriptionJob, job_id)
        if job:
            job.status = TranscriptionJobStatus.failed
            job.progress = 100
            job.error = str(exc)
            job.completed_at = datetime.utcnow()
            job.updated_at = datetime.utcnow()
            db.commit()
            log_event(db, actor_user_id, "transcription_failed", "transcript", job.transcript_id, {"job_id": job.id})
        logger.exception("Transcription job %s failed", job_id)
    finally:
        db.close()


@app.get("/api/transcripts", response_model=list[TranscriptOut])
def list_transcripts(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if user.access_level == AccessLevel.admin:
        transcripts = db.execute(select(Transcript)).scalars().all()
    else:
        transcripts = db.execute(select(Transcript).where(Transcript.owner_user_id == user.id)).scalars().all()
    return transcripts


@app.post("/api/transcripts", response_model=TranscriptOut)
def create_transcript(payload: TranscriptCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = Transcript(
        name=payload.name,
        description=payload.description,
        owner_user_id=user.id,
        status=TranscriptStatus.draft,
    )
    db.add(transcript)
    db.commit()
    db.refresh(transcript)
    log_event(db, user.id, "transcript_created", "transcript", transcript.id)
    return transcript


@app.get("/api/transcripts/{transcript_id}", response_model=TranscriptOut)

def get_transcript(transcript_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    return transcript


@app.patch("/api/transcripts/{transcript_id}", response_model=TranscriptOut)

def update_transcript(transcript_id: int, payload: TranscriptUpdate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    if payload.status == TranscriptStatus.completed:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Use mark-complete endpoint")
    data = payload.dict(exclude_unset=True)
    for key, value in data.items():
        setattr(transcript, key, value)
    db.commit()
    db.refresh(transcript)
    return transcript


@app.delete("/api/transcripts/{transcript_id}")
def delete_transcript(transcript_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    purge_transcript_audio(transcript.id)
    log_event(db, user.id, "audio_purged_for_delete", "transcript", transcript.id)
    db.delete(transcript)
    db.commit()
    log_event(db, user.id, "transcript_deleted", "transcript", transcript_id)
    return {"message": "deleted"}


@app.post("/api/transcripts/{transcript_id}/upload-wav")
def upload_wav(transcript_id: int, file: UploadFile = File(...), user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    if transcript.status == TranscriptStatus.completed:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Transcript completed")
    path = save_wav_file(transcript.id, file)
    transcript.wav_filename = sanitize_filename(file.filename)
    transcript.wav_storage_path = path
    db.commit()
    log_event(db, user.id, "wav_uploaded", "transcript", transcript.id)
    return {"message": "uploaded"}


@app.get("/api/transcripts/{transcript_id}/audio")
def get_transcript_audio(transcript_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    if transcript.status == TranscriptStatus.completed or not transcript.wav_storage_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No WAV available")
    if not os.path.exists(transcript.wav_storage_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="WAV file missing")
    return FileResponse(
        transcript.wav_storage_path,
        media_type="audio/wav",
        filename=transcript.wav_filename or "audio.wav",
    )


@app.post("/api/transcripts/{transcript_id}/transcribe", response_model=TranscriptionJobOut)
def transcribe(
    transcript_id: int,
    payload: StartTranscriptionRequest,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    transcript = get_transcript_or_404(transcript_id, user, db)
    if transcript.status == TranscriptStatus.completed:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Transcript completed")
    if not transcript.wav_storage_path:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="No WAV uploaded")
    hms_to_seconds(payload.start_time)

    existing_job = active_transcription_job(transcript.id, db)
    if existing_job:
        return existing_job

    job = TranscriptionJob(transcript_id=transcript.id, status=TranscriptionJobStatus.queued, progress=0)
    db.add(job)
    db.commit()
    db.refresh(job)
    log_event(db, user.id, "transcription_started", "transcript", transcript.id, {"job_id": job.id})
    background_tasks.add_task(run_transcription_job, job.id, payload.start_time, user.id)
    return job


@app.get("/api/transcripts/{transcript_id}/transcription-job", response_model=TranscriptionJobOut)
def get_transcription_job(transcript_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    job = latest_transcription_job(transcript.id, db)
    if not job:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No transcription job found")
    return job


@app.get("/api/transcripts/{transcript_id}/lines", response_model=list[TranscriptLineOut])
def list_lines(transcript_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    lines = db.execute(select(TranscriptLine).where(TranscriptLine.transcript_id == transcript.id).order_by(TranscriptLine.order_index)).scalars().all()
    return lines


@app.post("/api/transcripts/{transcript_id}/lines", response_model=TranscriptLineOut)
def add_line(transcript_id: int, payload: TranscriptLineCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    validate_speaker_label(transcript.id, payload.speaker_label_id, db)
    line = TranscriptLine(
        transcript_id=transcript.id,
        order_index=payload.order_index,
        timestamp_hms=payload.timestamp_hms,
        speaker_label_id=payload.speaker_label_id,
        text=payload.text,
    )
    db.add(line)
    db.commit()
    reorder_lines(transcript.id, db)
    db.refresh(line)
    return line


@app.patch("/api/transcripts/{transcript_id}/lines/{line_id}", response_model=TranscriptLineOut)
def update_line(transcript_id: int, line_id: int, payload: TranscriptLineUpdate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    line = db.get(TranscriptLine, line_id)
    if not line or line.transcript_id != transcript.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Line not found")
    data = payload.dict(exclude_unset=True)
    if "speaker_label_id" in data:
        validate_speaker_label(transcript.id, data["speaker_label_id"], db)
    for key, value in data.items():
        setattr(line, key, value)
    db.commit()
    reorder_lines(transcript.id, db)
    db.refresh(line)
    return line


@app.delete("/api/transcripts/{transcript_id}/lines/{line_id}")
def delete_line(transcript_id: int, line_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    line = db.get(TranscriptLine, line_id)
    if not line or line.transcript_id != transcript.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Line not found")
    db.delete(line)
    db.commit()
    reorder_lines(transcript.id, db)
    return {"message": "deleted"}


@app.post("/api/transcripts/{transcript_id}/lines/{line_id}/split")
def split_line(transcript_id: int, line_id: int, payload: SplitLineRequest, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    line = db.get(TranscriptLine, line_id)
    if not line or line.transcript_id != transcript.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Line not found")
    if payload.split_index <= 0 or payload.split_index >= len(line.text):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid split")
    first_text = line.text[: payload.split_index].strip()
    second_text = line.text[payload.split_index :].strip()
    line.text = first_text
    new_line = TranscriptLine(
        transcript_id=transcript.id,
        order_index=line.order_index + 1,
        timestamp_hms=line.timestamp_hms,
        speaker_label_id=line.speaker_label_id,
        text=second_text,
    )
    db.add(new_line)
    db.commit()
    reorder_lines(transcript.id, db)
    return {"message": "split"}


@app.post("/api/transcripts/{transcript_id}/lines/merge")
def merge_lines(transcript_id: int, payload: MergeLinesRequest, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    first_line = db.get(TranscriptLine, payload.first_line_id)
    second_line = db.get(TranscriptLine, payload.second_line_id)
    if not first_line or not second_line:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Line not found")
    if first_line.transcript_id != transcript.id or second_line.transcript_id != transcript.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Forbidden")
    if abs(first_line.order_index - second_line.order_index) != 1:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Lines must be adjacent")
    first_line.text = f"{first_line.text} {second_line.text}".strip()
    if payload.keep_label_id:
        first_line.speaker_label_id = payload.keep_label_id
    db.delete(second_line)
    db.commit()
    reorder_lines(transcript.id, db)
    return {"message": "merged"}


@app.post("/api/transcripts/{transcript_id}/labels", response_model=SpeakerLabelOut)
def create_label(transcript_id: int, payload: SpeakerLabelCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    label = SpeakerLabel(transcript_id=transcript.id, name=payload.name, color_hex=payload.color_hex)
    db.add(label)
    db.commit()
    db.refresh(label)
    return label


@app.get("/api/transcripts/{transcript_id}/labels", response_model=list[SpeakerLabelOut])
def list_labels(transcript_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    return db.execute(select(SpeakerLabel).where(SpeakerLabel.transcript_id == transcript.id)).scalars().all()


@app.patch("/api/transcripts/{transcript_id}/labels/{label_id}", response_model=SpeakerLabelOut)
def update_label(transcript_id: int, label_id: int, payload: SpeakerLabelUpdate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    label = db.get(SpeakerLabel, label_id)
    if not label or label.transcript_id != transcript.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Label not found")
    data = payload.dict(exclude_unset=True)
    for key, value in data.items():
        setattr(label, key, value)
    db.commit()
    db.refresh(label)
    return label


@app.get("/api/transcripts/{transcript_id}/vocabulary", response_model=list[EntryOut])
def list_vocab(transcript_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    return db.execute(select(VocabularyEntry).where(VocabularyEntry.transcript_id == transcript.id)).scalars().all()


@app.post("/api/transcripts/{transcript_id}/vocabulary", response_model=EntryOut)
def add_vocab(transcript_id: int, payload: EntryCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    entry = VocabularyEntry(owner_scope="transcript", transcript_id=transcript.id, word_or_phrase=payload.word_or_phrase)
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return entry


@app.delete("/api/transcripts/{transcript_id}/vocabulary/{entry_id}")
def delete_vocab(transcript_id: int, entry_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    entry = db.get(VocabularyEntry, entry_id)
    if not entry or entry.transcript_id != transcript.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Entry not found")
    db.delete(entry)
    db.commit()
    return {"message": "deleted"}


@app.get("/api/transcripts/{transcript_id}/exclude", response_model=list[EntryOut])
def list_exclude(transcript_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    return db.execute(select(ExcludeEntry).where(ExcludeEntry.transcript_id == transcript.id)).scalars().all()


@app.post("/api/transcripts/{transcript_id}/exclude", response_model=EntryOut)
def add_exclude(transcript_id: int, payload: EntryCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    entry = ExcludeEntry(owner_scope="transcript", transcript_id=transcript.id, word_or_phrase=payload.word_or_phrase)
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return entry


@app.delete("/api/transcripts/{transcript_id}/exclude/{entry_id}")
def delete_exclude(transcript_id: int, entry_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    entry = db.get(ExcludeEntry, entry_id)
    if not entry or entry.transcript_id != transcript.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Entry not found")
    db.delete(entry)
    db.commit()
    return {"message": "deleted"}


@app.post("/api/transcripts/{transcript_id}/export")
def export_transcript(transcript_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    lines = db.execute(select(TranscriptLine).where(TranscriptLine.transcript_id == transcript.id).order_by(TranscriptLine.order_index)).scalars().all()
    doc = Document()
    doc.add_heading(transcript.name, level=1)
    if transcript.description:
        doc.add_paragraph(transcript.description)
    doc.add_paragraph(f"Status: {transcript.status.value}")
    for line in lines:
        label = "Unknown"
        if line.speaker_label_id:
            label_obj = db.get(SpeakerLabel, line.speaker_label_id)
            if label_obj:
                label = label_obj.name
        doc.add_paragraph(f"[{line.timestamp_hms}] {label}: {line.text}")
    with tempfile.NamedTemporaryFile(prefix=f"atc_transcript_{transcript.id}_", suffix=".docx", delete=False) as export_file:
        export_path = export_file.name
    doc.save(export_path)
    log_event(db, user.id, "transcript_exported", "transcript", transcript.id)
    filename = f"{sanitize_filename(transcript.name).removesuffix('.wav')}_{transcript.id}.docx"
    return FileResponse(
        export_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
        background=BackgroundTask(lambda: os.path.exists(export_path) and os.remove(export_path)),
    )


@app.post("/api/transcripts/{transcript_id}/mark-complete")
def mark_complete(transcript_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    transcript = get_transcript_or_404(transcript_id, user, db)
    if transcript.status == TranscriptStatus.completed:
        return {"message": "already completed"}
    purge_transcript_audio(transcript.id)
    log_event(db, user.id, "audio_purged_on_complete", "transcript", transcript.id)
    transcript.status = TranscriptStatus.completed
    transcript.wav_storage_path = None
    transcript.wav_filename = None
    db.commit()
    log_event(db, user.id, "transcript_completed", "transcript", transcript.id)
    return {"message": "completed"}


@app.post("/api/admin/transcripts/{transcript_id}/reopen")
def reopen_transcript(transcript_id: int, admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    transcript = db.get(Transcript, transcript_id)
    if not transcript:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Transcript not found")
    transcript.status = TranscriptStatus.in_progress
    db.commit()
    log_event(db, admin.id, "transcript_reopened", "transcript", transcript.id)
    return {"message": "reopened"}
